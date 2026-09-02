from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
TARGET_DIR = BASE_DIR / "Nykyiset asiakkaat"
WORKBOOK = TARGET_DIR / "model_improvement_next_year_recent_weighted.xlsx"
MODEL_SCRIPT = BASE_DIR / "backtest_2025_model_improvements.py"
MODEL_SCRIPT_COPY = TARGET_DIR / "backtest_2025_model_improvements.py"
OUTPUT_DOCX = TARGET_DIR / "Innoflame_nykyasiakkaiden_potentiaalimalli_dokumentaatio.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
TEXT = "222222"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="000000")
        set_cell_shading(hdr[i], MID_GRAY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], "" if pd.isna(value) else value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    document.add_paragraph()
    return table


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def format_eur(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return ""
    return f"{float(value):,.0f} EUR".replace(",", " ")


def format_meur(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return ""
    return f"{float(value) / 1_000_000:.2f} MEUR"


def format_pct(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return ""
    return f"{float(value) * 100:.1f} %"


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Innoflame nykyasiakkaiden potentiaalimalli")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("000000")

    p = document.add_paragraph()
    run = p.add_run("Mallidokumentaatio: lähdetiedot, featuret, laskentalogiikka, tuoteryhmäkalibrointi ja outputit")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")

    meta = [
        ["Dokumentin tarkoitus", "Kuvata mitä tietoja malli tarvitsee, miten malli toimii ja miten tuloksia käytetään."],
        ["Pääoutput", str(WORKBOOK.name)],
        ["Malliskripti", str(MODEL_SCRIPT_COPY.name)],
        ["Rajaus", "Nykyiset asiakkaat, Innoflame poistettu asiakas- ja myyntihistorialähteistä."],
    ]
    add_table(document, ["Kenttä", "Arvo"], meta, [1.7, 4.8])


def load_workbook_context() -> dict[str, pd.DataFrame]:
    xl = pd.ExcelFile(WORKBOOK)
    return {sheet: pd.read_excel(WORKBOOK, sheet_name=sheet) for sheet in xl.sheet_names}


def add_executive_summary(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("1. Tiivistelmä", level=1)
    summary = sheets["summary"]
    latest = summary[summary["model"].eq("history_feature_model")].iloc[0]
    next_year = summary[summary["model"].eq("next_year_recent_weighted_forecast")].iloc[0]
    potential_case = summary[summary["model"].eq("sales_potential_case")].iloc[0]
    pg_raw = summary[summary["model"].eq("product_group_rows_raw")].iloc[0]
    pg_cal = summary[summary["model"].eq("product_group_rows_calibrated")].iloc[0]

    document.add_paragraph(
        "Malli arvioi nykyasiakkaiden vuositason potentiaalia ja tuoteryhmätason myyntimahdollisuuksia. "
        "Uusin versio käyttää asiakkaan ostohistoriaa, alkuperäisen prospektimallin signaaleja, todennäköisyyden kalibrointia ja tuoteryhmäkohtaista kalibrointikerrointa."
    )
    add_table(
        document,
        ["Mittari", "Arvo", "Tulkinta"],
        [
            ["2025 toteuma", format_meur(latest["actual_sales_2025_eur"]), "Vertailutoteuma backtestille."],
            ["Uusimman mallin ennuste", format_meur(latest["predicted_eur"]), "Ostohistoriafeatureihin perustuva vuosiarvio."],
            ["Vuoden 2027 run-rate", format_meur(next_year["predicted_eur"]), "Konservatiivinen vuoden 2027 ennuste nykyisella ostotasolla ja 2026 YTD annualisoinnilla."],
            ["Myynnillinen potentiaalicase", format_meur(potential_case["predicted_eur"]), "Run-rate + kalibroiduista tuoteryhmista laskettu kasvumahdollisuus."],
            ["Potentiaalicasen kasvu 2025 vs.", format_pct(potential_case["mape_actual_over_100eur"]), "Kuinka paljon realistinen potentiaali ylittaa vuoden 2025 toteuman."],
            ["Kokonaisbias", format_meur(latest["bias_eur"]), "Ennusteen kokonaisero toteumaan."],
            ["Korrelaatio", f"{latest['correlation']:.2f}", "Asiakaskohtaisen arvion yhteys toteumaan."],
            ["Growth AUC", f"{latest['growth_auc']:.2f}", "Kasvutodennäköisyyden erottelukyky."],
            ["Tuoteryhmämalli, raaka", format_meur(pg_raw["predicted_eur"]), "White space -malli ennen kalibrointia."],
            ["Tuoteryhmämalli, kalibroitu", format_meur(pg_cal["predicted_eur"]), "Tuoteryhmäkohtaisilla kertoimilla korjattu arvo."],
        ],
        [2.1, 1.4, 3.0],
    )


def add_input_data_section(document: Document) -> None:
    document.add_heading("2. Mallin tarvitsemat lähtötiedot", level=1)
    document.add_paragraph("Malli tarvitsee seuraavat lähdetiedostot. Polut ovat projektikansion suhteellisia tiedostoja.")
    add_table(
        document,
        ["Tiedosto", "Rooli mallissa", "Keskeiset kentät"],
        [
            [
                "GoSystems_accounts_25_06_2026_updated_business_ids_fi_normalized_without_innoflame.xlsx",
                "Nykyasiakkaiden perusjoukko ja tunnisteet.",
                "id/account_id, name, company_name, business_id, country, category",
            ],
            [
                "GoSystems_sales_26_05_2026_summarized_without_innoflame.csv",
                "Ostohistoria ja 2025 backtest-toteuma.",
                "account_id, sku, category, reference, created_year_month, total_value",
            ],
            [
                "current_customer_potential_without_innoflame.csv",
                "Alkuperäisestä prospektimallista johdetut score- ja potentiaalisignaalit.",
                "score, probability_of_growth, expected_potential_eur, company_segment, positive_signals",
            ],
            [
                "product_group_recommendations_without_innoflame.csv",
                "Alkuperäisen tuoteryhmäsuosituksen lähtöarvot.",
                "product_group_code, product_group_name, white_space_gap, recommended_group_expected_potential_eur",
            ],
            [
                "product_master_enrichment/final_product_grouping/Innoflame_tuoteryhmittely.xlsx",
                "SKU-rivien mapitus alimman saatavilla olevan tuoteryhmätason mukaan.",
                "sku/code/product_id, product_group_l1-l4_code, product_group_l1-l4_name",
            ],
        ],
        [2.2, 2.0, 2.3],
    )
    document.add_paragraph(
        "Y-tunnus normalisoidaan muotoon 1234567-8. Jos tunniste alkaa FI-etuliitteellä, etuliite poistetaan ja väliviiva asetetaan ennen viimeistä numeroa."
    )


def add_feature_section(document: Document) -> None:
    document.add_heading("3. Mallin käyttämät featuret", level=1)
    document.add_paragraph("Featuret muodostetaan vain tiedoista, jotka olisivat tiedossa ennen vuoden 2025 toteuman arviointia. Vuoden 2025 myyntiä käytetään backtestin targetina.")
    add_table(
        document,
        ["Feature-ryhmä", "Sarakkeet", "Miksi mukana"],
        [
            ["Asiakastunnisteet", "business_id, account_id, customer_name, country, category", "Yhdistäminen, riviseuranta ja raportointi."],
            ["Historiallinen myynti", "sales_2023_eur, sales_2024_eur, sales_2024_q1-q4_eur", "Vuositason myyntikoko, kausivaihtelu ja pohjataso."],
            ["Aktiivisuus", "active_months_2023_2024, active_months_2024, order_rows_2024", "Kuinka jatkuva ja toistuva asiakassuhde on."],
            ["Momentum", "sales_momentum_2024_vs_2023, h2_vs_h1_2024", "Onko ostaminen kasvussa vai laskussa ennen ennustevuotta."],
            ["Recency", "days_since_last_purchase_at_2025_start", "Kuinka tuore ostosuhde oli vuoden 2025 alussa."],
            ["Tuoteryhmälaajuus", "product_group_count_2024", "Kuinka monipuolisesti asiakas ostaa eri ryhmiä."],
            ["Prospektimallin signaalit", "score, probability_of_growth, revenue_k_eur, segment_lift", "Yritys- ja segmenttitason signaalit alkuperäisestä malliputkesta."],
        ],
        [1.6, 2.4, 2.5],
    )


def add_model_logic_section(document: Document) -> None:
    document.add_heading("4. Itse malli ja laskentalogiikka", level=1)
    document.add_paragraph("Malli koostuu asiakastason ennusteesta, kasvutodennäköisyydestä, tuoteryhmämallista ja myynnin palautesilmukasta.")
    add_table(
        document,
        ["Vaihe", "Menetelmä", "Output"],
        [
            ["1. Datan luku ja normalisointi", "Lähtötiedostot luetaan, business_id normalisoidaan ja myyntirivit yhdistetään asiakkaisiin account_id:n kautta.", "Yhtenäinen asiakas- ja myyntidata."],
            ["2. Ostohistoriafeaturet", "Vuosi-, kvartaali-, aktiivisuus-, momentum- ja recency-featuret lasketaan 2023-2024 datasta.", "history_features."],
            ["3. Todennäköisyysmalli", "RandomForestClassifier arvioi grew_2025-targetin. Todennäköisyys kalibroidaan myös isotonic-regressiolla.", "improved_probability_of_growth, calibrated_current_probability."],
            ["4. Vuositason euromalli", "RandomForestRegressor ennustaa vuoden 2025 euromääräistä myyntiä ostohistoria- ja score-featureilla.", "improved_expected_sales_2025_eur."],
            ["5. Tuoteryhmämalli", "Asiakas x tuoteryhmä -tasolla lasketaan customer_share, similar_customer_share ja white_space_gap.", "product_group_model_expected_2025_eur."],
            ["6. Tuoteryhmäkalibrointi", "Toteuma/ennuste-suhteesta johdetaan tuoteryhmäkohtainen kerroin, jota tasoitetaan globaalilla kertoimella pienissä ryhmissä.", "product_group_calibrated_expected_2025_eur ja recommendations_calibrated."],
            ["7. CRM-potentials validointi", "Status, Sales ja Probability aggregoidaan asiakastasolle. Jos CRM-osumaa ei löydy, alkuperäinen malliarvo säilyy.", "crm_potential_validation."],
            ["8. Virheanalyysi", "Mallin ennustetta verrataan 2025 toteumaan ja asiakkaat luokitellaan virhebuckettiin.", "error_analysis ja sales_feedback_template."],
        ],
        [1.5, 3.2, 1.8],
    )
    document.add_paragraph("Mallin tekninen toteutus on tallennettu samaan kansioon tiedostoon backtest_2025_model_improvements.py.")


def add_product_group_section(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("5. Tuoteryhmälogiikka ja kalibrointi", level=1)
    document.add_paragraph(
        "Tuoteryhmäsuositus ei ehdota yksittäisiä SKU-tuotteita. SKU-rivit mapataan ensisijaisesti tuoteryhmittelyn alimpaan saatavilla olevaan tasoon. "
        "Jos SKU ei löydy tuotemasterista, käytetään myyntirivin category/reference-tietoa fallback-ryhmänä."
    )
    pg = sheets["product_group_calibration"].head(8).copy()
    add_table(
        document,
        ["Tuoteryhmä", "2025 toteuma", "Raaka ennuste", "Kalibroitu ennuste", "Kerroin"],
        [
            [
                row["lowest_product_group_name"],
                format_eur(row["actual_group_sales_2025_eur"]),
                format_eur(row["product_group_model_expected_2025_eur"]),
                format_eur(row["product_group_calibrated_expected_2025_eur"]),
                f"{row['product_group_calibration_factor']:.2f}",
            ]
            for _, row in pg.iterrows()
        ],
        [2.0, 1.2, 1.2, 1.3, 0.8],
    )
    add_bullets(
        document,
        [
            "Raaka tuoteryhmämalli perustuu white space -eroon asiakkaan oman ostojakauman ja saman segmentin ostojakauman välillä.",
            "Kalibrointikerroin korjaa tuoteryhmäkohtaisen aliarvion tai yliarvion vuoden 2025 backtestin perusteella.",
            "Isoissa ryhmissä käytetään ryhmän omaa kerrointa. Pienissä ryhmissä kerrointa tasoitetaan globaalilla toteuma/ennuste-suhteella.",
            "Kalibroitu suositus löytyy välilehdeltä recommendations_calibrated.",
        ],
    )


def add_sales_potential_section(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("6. Myynnillinen potentiaalicase", level=1)
    summary = sheets["sales_potential_summary"].set_index("metric")["value"]
    sample = sheets["sales_potential_case"].head(10).copy()

    document.add_paragraph(
        "Uusin tulos erottaa kaksi eri kayttotarkoitusta: konservatiivinen run-rate ennuste kertoo, "
        "mihin myynti todennakoisesti asettuu nykyisella ostotasolla, kun taas myynnillinen potentiaalicase "
        "nostaa esiin asiakkaat ja tuoteryhmat, joissa aktiivisella myyntityolla voidaan hakea kasvua."
    )
    add_table(
        document,
        ["Mittari", "Arvo", "Tulkinta"],
        [
            ["2025 toteuma", format_meur(summary.get("actual_sales_2025_eur")), "Historiallinen vertailutaso."],
            ["Vuoden 2027 run-rate ennuste", format_meur(summary.get("base_forecast_2027_eur")), "Konservatiivinen vuoden 2027 ennuste samalla laskennalla."],
            ["Tuoteryhmien bruttokasvupooli", format_meur(summary.get("product_group_growth_pool_eur")), "Kalibroitujen tuoteryhmasuositusten summa ennen asiakaskohtaista rajausta."],
            ["Oikaistu kasvupotentiaali", format_meur(summary.get("growth_potential_eur")), "Todennakoisyys-, prioriteetti- ja katto-oikaistu kasvu."],
            ["Realistinen 2027 potentiaali", format_meur(summary.get("realistic_potential_2027_eur")), "Vertailutaso + oikaistu kasvupotentiaali."],
            ["Upside 2027 potentiaali", format_meur(summary.get("upside_potential_2027_eur")), "Korkeampi skenaario tuoteryhma-avauksille."],
            ["Kasvu 2025 toteumaan", format_pct(summary.get("realistic_potential_vs_2025_pct")), "Realistisen potentiaalin ero suhteessa 2025 toteumaan."],
        ],
        [2.1, 1.4, 3.0],
    )
    document.add_paragraph(
        "Asiakaskohtaisessa sales_potential_case-valilehdessa keskeiset sarakkeet ovat base_forecast_eur, "
        "growth_potential_eur, realistic_potential_eur, upside_potential_eur, improved_probability_of_growth "
        "ja top_recommended_product_groups."
    )
    add_table(
        document,
        ["Asiakas", "2025 toteuma", "Run-rate", "Realistinen potentiaali", "Suositellut tuoteryhmat"],
        [
            [
                row.get("customer_name") or row.get("company"),
                format_eur(row.get("actual_sales_2025_eur")),
                format_eur(row.get("base_forecast_eur")),
                format_eur(row.get("realistic_potential_eur")),
                row.get("top_recommended_product_groups", ""),
            ]
            for _, row in sample.iterrows()
        ],
        [1.8, 1.1, 1.1, 1.2, 1.8],
    )


def add_crm_validation_section(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("7. CRM-potentials validointi", level=1)
    summary = sheets["crm_validation_summary"].set_index("metric")["value"]
    sample = sheets["crm_potential_validation"].head(8).copy()

    document.add_paragraph(
        "CRM-potentials-aineistoa kaytetaan mallin realistisuuden tarkistamiseen. CRM-riveilla voi olla samalle "
        "yritykselle useita mahdollisuuksia, joten rivit aggregoidaan asiakastasolle nimen perusteella. "
        "CRM-odotusarvo lasketaan kaavalla Sales * Probability."
    )
    add_table(
        document,
        ["Mittari", "Arvo", "Tulkinta"],
        [
            ["CRM-riveja", f"{int(summary.get('crm_rows_input', 0)):,}".replace(",", " "), "Alkuperaiset CRM-potentials rivit."],
            ["CRM-nimia", f"{int(summary.get('crm_unique_names', 0)):,}".replace(",", " "), "CRM-rivit aggregoituna normalisoidulle nimelle."],
            ["CRM-osuma malliasiakkaalle", f"{int(summary.get('model_customers_with_crm_match', 0)):,}".replace(",", " "), "Malliasiakkaat, joille loytyi CRM-osuma."],
            ["Ei CRM-osumaa", f"{int(summary.get('model_customers_without_crm_match_kept_original', 0)):,}".replace(",", " "), "Nailla asiakkailla alkuperainen malliarvo sailyi."],
            ["Mahdollisesti liian pieni", f"{int(summary.get('model_may_be_too_low_count', 0)):,}".replace(",", " "), "CRM expected sales yli 20 % realistisen potentiaalin."],
            ["Nostettu CRM-odotusarvoon", f"{int(summary.get('raised_to_crm_expected_sales_count', 0)):,}".replace(",", " "), "Validointiarvo nostettiin CRM expected sales -tasolle."],
        ],
        [2.2, 1.2, 3.1],
    )
    document.add_paragraph(
        "Tarkistussaanto on varovainen: jos CRM-osumaa ei loydy, crm_validated_realistic_potential_2027_eur "
        "sailyttaa alkuperaisen realistic_potential_2027_eur-arvon. CRM ei siis nollaa tai pienennna mallin potentiaalia puuttuvan osuman takia."
    )
    add_table(
        document,
        ["Asiakas", "Malli 2027", "CRM expected", "Validointiarvo", "Status"],
        [
            [
                row.get("customer_name") or row.get("company"),
                format_eur(row.get("realistic_potential_2027_eur")),
                format_eur(row.get("crm_expected_sales_eur")),
                format_eur(row.get("crm_validated_realistic_potential_2027_eur")),
                row.get("crm_validation_status", ""),
            ]
            for _, row in sample.iterrows()
        ],
        [1.8, 1.1, 1.1, 1.2, 1.3],
    )


def add_output_section(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("8. Outputit ja välilehtien käyttötarkoitus", level=1)
    guide = [
        ["summary", "Mallin päätason mittarit.", "Aloita tästä: toteuma, ennuste, bias, korrelaatio ja AUC."],
        ["next_year_forecast", "Seuraavan vuoden konservatiivinen run-rate ennuste.", "Kayta perusennusteena nykyisella ostotasolla."],
        ["next_year_summary", "Run-rate ennusteen yhteenveto.", "Tarkista 2026 YTD annualisointi ja ero vuoden 2025 toteumaan."],
        ["sales_potential_case", "Myynnillinen potentiaalicase asiakastasolla.", "Priorisoi asiakkaat, joilla on suurin realistinen kasvumahdollisuus."],
        ["sales_potential_summary", "Potentiaalicasen yhteenveto.", "Vertaa run-rate ennustetta ja myynnillista potentiaalia."],
        ["crm_potential_validation", "CRM Status/Sales/Probability -validointi.", "Tarkista onko malli liian pieni tai linjassa CRM-pipelinen kanssa."],
        ["crm_validation_summary", "CRM-validoinnin yhteenveto.", "Nayttaa osumat, puuttuvat CRM-osumat ja nostetut validointiarvot."],
        ["crm_unmatched_names", "CRM-nimet ilman account-osumaa.", "Kayta nimikohdistuksen ja asiakasrekisterin laadun parantamiseen."],
        ["customer_backtest_2025", "Asiakaskohtainen backtest.", "Etsi asiakkaat, joissa malli osuu tai poikkeaa toteumasta."],
        ["history_features", "Ostohistoriafeaturet.", "Tarkista mistä asiakkaan vuosiarvio muodostuu."],
        ["probability_calibration", "Todennäköisyyden kalibrointi.", "Arvioi probability_of_growth-sarakkeen realistisuutta."],
        ["product_group_model", "Asiakas x tuoteryhmä -tason laskenta.", "Tutki white space -arviota ja tuoteryhmätoteumaa."],
        ["product_group_summary", "Tuoteryhmien yhteenveto.", "Vertaa raakaa ja kalibroitua tuoteryhmäennustetta."],
        ["product_group_calibration", "Tuoteryhmäkertoimet.", "Selittää miksi ryhmät, kuten Sales promotion tai työvaatetus, skaalautuvat."],
        ["recommendations_calibrated", "Kalibroitu tuoteryhmäsuositus.", "Käytä myynnin asiakaskohtaiseen tuoteryhmäpriorisointiin."],
        ["error_analysis", "Virhebucketit ja suurimmat poikkeamat.", "Valitse tarkistettavat asiakkaat."],
        ["sales_feedback_template", "Myynnin palautepohja.", "Täytä korjattu potentiaali, syy ja kommentti seuraavaa mallia varten."],
        ["feature_importance", "Mallin tärkeimmät signaalit.", "Ymmärrä mitkä featuret vaikuttavat eniten."],
        ["model_notes", "Mallin dokumentaatiorivit.", "Tarkista mallin rajaukset ja ajon periaatteet."],
    ]
    add_table(document, ["Välilehti", "Miksi mukana", "Käyttö"], guide, [1.8, 2.0, 2.7])


def add_validation_section(document: Document, sheets: dict[str, pd.DataFrame]) -> None:
    document.add_heading("9. Validointi ja laadunvarmistus", level=1)
    summary = sheets["summary"]
    latest = summary[summary["model"].eq("history_feature_model")].iloc[0]
    error_counts = sheets["error_analysis"]["error_bucket"].value_counts().reset_index()
    error_counts.columns = ["error_bucket", "customers"]
    add_bullets(
        document,
        [
            f"Asiakaskohtaisen vuosimyyntiarvion korrelaatio 2025 toteumaan on {latest['correlation']:.2f}.",
            f"Kasvutodennäköisyyden AUC on {latest['growth_auc']:.2f}.",
            f"Kokonaisbias uusimmassa mallissa on {format_meur(latest['bias_eur'])}.",
            "Tuoteryhmäkalibrointi nostaa top 60 -ryhmien ennusteen raakaversiosta lähelle vuoden 2025 toteumaa.",
            "Virheanalyysi säilytetään, jotta myynti voi antaa mallille korjaukset takaisin seuraavaan ajoon.",
        ],
    )
    add_table(
        document,
        ["Virhebucket", "Asiakkaita", "Tulkinta"],
        [
            [row["error_bucket"], int(row["customers"]), bucket_description(row["error_bucket"])]
            for _, row in error_counts.iterrows()
        ],
        [1.8, 1.1, 3.6],
    )


def bucket_description(bucket: str) -> str:
    return {
        "good_fit": "Malli on riittävän lähellä toteumaa.",
        "medium_error": "Poikkeama on keskisuuri ja kannattaa seurata.",
        "model_over_high": "Malli arvioi asiakkaan selvästi toteumaa korkeammaksi.",
        "model_under_high": "Toteuma on selvästi malliarviota korkeampi.",
    }.get(bucket, "Tarkistettava virheluokka.")


def add_limitations_section(document: Document) -> None:
    document.add_heading("10. Rajaukset ja seuraavat parannukset", level=1)
    add_bullets(
        document,
        [
            "Backtest perustuu vuoden 2025 toteumaan. Malli kannattaa validoida uudelleen seuraavalla täydellä toteumavuodella.",
            "Tuoteryhmän fallback category/reference-tasolle on käytännöllinen, mutta ei yhtä tarkka kuin SKU:n täysi tuotemaster-osuma.",
            "Todennäköisyys ja euromääräinen vuosiarvio ovat eri asioita. Niitä ei pidä tulkita samaksi mittariksi.",
            "Myynnin palaute kannattaa liittää takaisin opetusdataan: korjattu potentiaali, poikkeussyy, poissuljettavat asiakkaat ja puuttuvat tuoteryhmät.",
            "Malli ei korvaa asiakasvastuullisen arviota, vaan antaa priorisoidun työlistan ja perustellun lähtöpisteen.",
        ],
    )


def add_reproducibility_section(document: Document) -> None:
    document.add_heading("11. Ajon toistaminen", level=1)
    document.add_paragraph("Ajo voidaan toistaa projektikansiossa seuraavalla skriptillä:")
    p = document.add_paragraph()
    run = p.add_run("python backtest_2025_model_improvements.py")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    document.add_paragraph("Skripti tuottaa Excelin model_improvement_next_year_recent_weighted.xlsx ja CSV-tiedostot kansioon model_improvement_backtest_2025.")
    add_table(
        document,
        ["Artefakti", "Sijainti"],
        [
            ["Mallidokumentaatio", str(OUTPUT_DOCX.name)],
            ["Malliskripti", str(MODEL_SCRIPT_COPY.name)],
            ["Pääoutput Excel", str(WORKBOOK.name)],
        ],
        [2.0, 4.5],
    )


def main() -> None:
    TARGET_DIR.mkdir(exist_ok=True)
    sheets = load_workbook_context()

    document = Document()
    configure_styles(document)
    add_title(document)
    add_executive_summary(document, sheets)
    add_input_data_section(document)
    add_feature_section(document)
    add_model_logic_section(document)
    add_product_group_section(document, sheets)
    add_sales_potential_section(document, sheets)
    add_crm_validation_section(document, sheets)
    add_output_section(document, sheets)
    add_validation_section(document, sheets)
    add_limitations_section(document)
    add_reproducibility_section(document)

    document.core_properties.title = "Innoflame nykyasiakkaiden potentiaalimalli"
    document.core_properties.subject = "Mallidokumentaatio"
    document.core_properties.author = "OpenAI Codex"
    document.save(OUTPUT_DOCX)

    shutil.copy2(MODEL_SCRIPT, MODEL_SCRIPT_COPY)
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Copied model script to {MODEL_SCRIPT_COPY}")


if __name__ == "__main__":
    main()
