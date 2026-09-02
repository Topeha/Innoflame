from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "Innoflame_tuoteryhmittely_menetelmadokumentti.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_borders(table, color="C9D3E3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_text(cell, text)
    set_table_width(table, [6.5])
    doc.add_paragraph()


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Innoflame tuoteryhmittely - menetelmäkuvaus")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Innoflamen tuoteryhmittelyn rakentaminen")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Menetelmäkuvaus ja käytetyt lähteet")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_callout(
        doc,
        "Tuoteryhmittely on rakennettu päätuotetasolle. Lopullisessa aineistossa on 11 381 päätuotetta, "
        "14 päätasoa, 54 alatason ryhmää ja 109 tarkinta ryhmää. Variantit ovat mukana tuotemasterin "
        "rakenteessa, mutta niitä ei lasketa erillisinä tuotteina tuoteryhmäpuun yhteenvedossa.",
    )

    doc.add_heading("1. Lopputulos", level=1)
    doc.add_paragraph(
        "Työn lopputuloksena syntyi Innoflamelle kolmitasoinen tuoteryhmäpuu, jossa jokaisella päätuotteella "
        "on päätaso, alatason ryhmä ja tarkin tuoteryhmä. Neljäs taso jätettiin pois käytännön selkeyden vuoksi, "
        "koska se toisti useissa kohdissa ylempiä tasoja tai johti liian pieniin ryhmiin."
    )
    add_table(
        doc,
        ["Mittari", "Arvo", "Selite"],
        [
            ["Päätuotteet", "11 381", "Yksi rivi vastaa päätuotetta, esimerkiksi IF-9271."],
            ["Variantit tuotemasterissa", "34 552", "Variantit löytyvät products.json.gz-tiedoston options-rakenteesta."],
            ["Päätasot", "14", "Tuoteryhmäpuun ylin taso."],
            ["Taso 2", "54", "Päätason alla oleva ryhmittely."],
            ["Taso 3", "109", "Tarkin käytössä oleva ryhmä tässä versiossa."],
            ["Minimikoko alimmalla tasolla", "5", "Pienemmät ryhmät yhdistettiin tai ohjattiin tarkistettaviin."],
        ],
        [1.8, 1.0, 3.7],
    )

    doc.add_heading("2. Käytetyt lähteet", level=1)
    doc.add_paragraph(
        "Tuoteryhmittely perustuu Innoflamen tuotemasteriin ja sitä rikastaviin lähteisiin. Pääasiallisena "
        "tuotelähteenä käytettiin tuotemasteria, ja lisätiedot haettiin tuotteen omista teksteistä, toimittaja-aineistoista "
        "sekä brändi- ja varastotiedoista."
    )
    add_table(
        doc,
        ["Lähde", "Käyttötapa"],
        [
            ["products.json.gz", "Päätuotteet, varianttirakenne, tuotekoodit, tuotetekstit ja GoSystem-tuotemasterin perusrakenne."],
            ["products_table_view.csv", "Tuotemasterista muodostettu taulukkonäkymä analyysiin ja välivaiheisiin."],
            ["Innoflame_tuoteryhmittely.csv/xlsx", "Lopullinen tuoteryhmittely ja päätuotetason yhteenveto."],
            ["Brand mapping.xlsx", "Bränditietojen täydentäminen ja yhtenäistäminen."],
            ["Product lists from suppliers.zip", "Toimittajakohtaiset lisätiedot, kuten paino-, koko- ja logistiikkatietojen rikastus."],
            ["GC_tuotetiedot.xlsx", "GC-tuotteiden ja pakkaus-/laatikkomittojen täydennys soveltuvin osin."],
            ["Trexet_SS26.xlsx", "Trexet/New Wave -tuotteiden lisätiedot."],
            ["products_Stanley_Stella.csv", "Stanley/Stella-tuotteiden paino- ja tuotetietojen täydennys."],
            ["Fiskars-aineistot", "Fiskarsin tuote- ja logistiikkatietojen rikastus."],
            ["Asiakaspalaute tuoteryhmäpuusta", "Tason 4 poistaminen pääosin, nimistömuutokset, promootio-/decal-/lahjakorttiryhmien tarkennukset."],
        ],
        [2.3, 4.2],
    )

    doc.add_heading("3. Kentät, joilla luokittelua ohjattiin", level=1)
    doc.add_paragraph(
        "Luokittelu tehtiin ensisijaisesti tuotteen omien tekstien ja tunnisteiden perusteella. Erityisesti hyödynnettiin "
        "tuotenimeä, suomenkielistä otsikkoa, suomenkielistä kuvausta, hakudataa, SKU-/tuotekoodia ja brändinimeä."
    )
    add_bullet(doc, "Tuotteen tunnisteet: product_id, code, sku.")
    add_bullet(doc, "Tuotteen tekstisisältö: product_name, title_fi, description_fi ja searchdata.")
    add_bullet(doc, "Täydentävät tiedot: brand_name, inventory_category ja inventory_supplier.")
    add_bullet(doc, "Lopullista uutta luokittelua ei ohjattu inventory_warehouse_category-kentällä, koska se haluttiin rajata pois ohjaavana tietona.")

    doc.add_heading("4. Rakentamisen vaiheet", level=1)
    add_number(doc, "Tuotemasteri purettiin päätuotetasolle. Variantit tunnistettiin products.json.gz-tiedoston options-listasta, mutta tuoteryhmittely tehtiin päätuotteille.")
    add_number(doc, "Tuotetekstit normalisoitiin hakua varten. Luokittelussa yhdistettiin tuotenimi, otsikko, kuvaus, hakudata, koodi ja brändi.")
    add_number(doc, "Tuotteille luotiin ensimmäinen tuoteryhmäpuu olemassa olevien tuotekategorioiden, tekstisääntöjen ja tuotetyyppien perusteella.")
    add_number(doc, "Toimittaja- ja brändiaineistoilla rikastettiin masteria, jotta tuotetekstit, mitat ja tunnisteet tukivat myöhempää laadunvarmistusta.")
    add_number(doc, "inventory_warehouse_category poistettiin ohjaavasta roolista ja luokittelu rakennettiin uudelleen tuotetekstien, tunnisteiden ja sääntöjen pohjalta.")
    add_number(doc, "Liian pienet alimman tason ryhmät yhdistettiin, jotta alimmalla tasolla ei jäisi alle viiden tuotteen ryhmiä.")
    add_number(doc, "Asiakaspalautteen perusteella ryhmäpuu yksinkertaistettiin pääosin kolmeen tasoon ja nimistöä korjattiin.")
    add_number(doc, "Tarkistettavat-ryhmälle tehtiin erillinen analyysi. Korkean varmuuden ehdotukset siirrettiin automaattisesti oikeisiin ryhmiin.")
    add_number(doc, "Lopuksi muodostettiin Excel, PowerPoint-kooste ja päätuotetason yhteenvedot.")

    doc.add_heading("5. Asiakaspalautteen perusteella tehdyt keskeiset muutokset", level=1)
    add_bullet(doc, "Taso 4 poistettiin käytöstä pääosin, koska se teki puusta liian spesifin ja toisti usein tason 3 otsikoita.")
    add_bullet(doc, "Nimistöä yhtenäistettiin niin, ettei sama otsikko toistu samalla polulla tasolla 2 ja tasolla 3.")
    add_bullet(doc, "Kaulanauhat, avaimenperät, heijastimet, pinssit ja rintanapit sijoitettiin promootio- ja tapahtumatuotteiden alle.")
    add_bullet(doc, "Putkihuivit erotettiin tavallisista kaulahuiveista ja sijoitettiin promootio-/jakotuotteisiin.")
    add_bullet(doc, "Decalit ja merkkaustuotteet muodostettiin omaksi kokonaisuudekseen siirtokuville, tekstiilimerkeille ja vastaaville tuotteille.")
    add_bullet(doc, "Lahjat ja sesonkituotteet korvattiin rajatummalla lahjakortit ja hyväntekeväisyys -kokonaisuudella.")
    add_bullet(doc, "Asusteet sijoitettiin vaatteiden alle, ja korut, kellot sekä aurinkolasit pidettiin omana päätasonaan.")
    add_bullet(doc, "Turvallisuus-käsitettä kavennettiin; suojaimet säilyivät omana käyttötarkoitusta paremmin kuvaavana kokonaisuutena.")

    doc.add_heading("6. Tarkistettavat-ryhmä ja automaattiset siirrot", level=1)
    doc.add_paragraph(
        "Tarkistettavat-ryhmä kerää tuotteet, joiden tekstistä tai tunnisteista ei saatu riittävän varmaa kohdistusta. "
        "Ryhmästä tehtiin erillinen analyysi, jossa haettiin mahdollisia siirtoja muihin tuoteryhmiin avainsanojen ja "
        "kuvaustekstien perusteella."
    )
    add_table(
        doc,
        ["Kohta", "Määrä", "Selite"],
        [
            ["Tarkistettavat tuotteet", "1 033", "Noin 9,1 % kaikista päätuotteista analyysin jälkeen."],
            ["Korkean varmuuden siirrot", "117", "Siirrettiin automaattisesti ehdotettuun ryhmään."],
            ["Kohderyhmiä siirroissa", "10", "High-ehdotukset jakautuivat kymmeneen tuoteryhmäpolkuun."],
            ["Manuaalisesti tarkistettavat", "906", "Ei riittävän vahvaa tekstiosumaa automaattiseen siirtoon."],
        ],
        [2.1, 1.0, 3.4],
    )

    doc.add_heading("7. Laadunvarmistus", level=1)
    add_bullet(doc, "Päätuotemäärä tarkistettiin product_id-tasolla: 11 381 uniikkia päätuotetta.")
    add_bullet(doc, "Variantit tunnistettiin erillisestä options-rakenteesta, eikä niitä laskettu tuoteryhmäpuun tuotteiksi.")
    add_bullet(doc, "Alimman tason minimikooksi asetettiin viisi tuotetta, jotta yksittäistuotteiden ryhmiä ei jäisi käyttöön.")
    add_bullet(doc, "Taso 2 ja taso 3 tarkistettiin niin, ettei sama otsikko toistu samassa polussa epäloogisesti.")
    add_bullet(doc, "Lopullisen Excelin Yhteenveto-välilehti päivitettiin laskemaan päätuotteita, ei variantteja.")
    add_bullet(doc, "PowerPoint-kooste päivitettiin käyttämään lähteenä Innoflame_tuoteryhmittely.csv-tiedostoa.")

    doc.add_heading("8. Tuoteryhmittelyn tiedostot", level=1)
    add_table(
        doc,
        ["Tiedosto", "Sisältö"],
        [
            ["product_master_enrichment/final_product_grouping/Innoflame_tuoteryhmittely.xlsx", "Varsinainen Excel-tuoteryhmittely ja päätuotetason yhteenveto."],
            ["product_master_enrichment/final_product_grouping/Innoflame_tuoteryhmittely.csv", "Sama tuoteryhmittely CSV-muodossa jatkokäsittelyä varten."],
            ["outputs/Innoflame_tuoteryhmittely_kooste_paivitetty.pptx", "PowerPoint-kooste tuoteryhmäpuun rakenteesta ja keskeisistä havainnoista."],
            ["outputs/product_grouping_summary/paatotuotteet_ryhmittain_yhteenveto.xlsx", "Erillinen päätuotemäärien yhteenveto ryhmittäin."],
            ["outputs/product_grouping_summary/tarkistettavat_ryhma_analyysi.xlsx", "Tarkistettavat-ryhmän analyysi ja jatkokohdistusehdotukset."],
        ],
        [3.4, 3.1],
    )

    doc.add_heading("9. Rajaukset ja jatkohuomiot", level=1)
    add_bullet(doc, "Tuoteryhmittely on tarkoitettu tuotepuun ja raportoinnin rakenteeksi päätuotetasolla.")
    add_bullet(doc, "Variantit perivät päätuotteen tuoteryhmän, koska varianttiriveillä ei ole omaa tuoteryhmäkenttää tuotemasterissa.")
    add_bullet(doc, "Tarkistettavat-ryhmään jääneet tuotteet vaativat sisällöllistä läpikäyntiä, jos niille halutaan tarkempi kohdistus.")
    add_bullet(doc, "Uusien tuotteiden luokittelussa kannattaa ylläpitää samoja sääntöjä ja täydentää avainsanoja niistä tapauksista, jotka nyt jäivät manuaaliseen tarkistukseen.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
