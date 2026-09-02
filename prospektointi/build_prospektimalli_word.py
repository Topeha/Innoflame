from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "PROSPEKTIMALLI_TAULUKUVAUS.md"
OUTPUT = ROOT / "Prospektimallin_taulukuvaus.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
BORDER = "B8C4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown_text(paragraph, text: str, *, size=None, bold_default=False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        clean = part[1:-1] if is_code else part
        run = paragraph.add_run(clean)
        set_font(run, size=size, bold=bold_default or is_code)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    body = rows[2:] if all(set(c.strip()) <= {"-", ":"} for c in rows[1]) else rows[1:]
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if col_count == 2:
        widths = [2400, 6960]
    elif col_count == 3:
        widths = [2400, 3000, 3960]
    else:
        widths = [int(9360 / col_count)] * col_count

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_markdown_text(p, text.strip(), size=9, bold_default=True)

    for row in body:
        cells = table.add_row().cells
        for idx in range(col_count):
            text = row[idx].strip() if idx < len(row) else ""
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_markdown_text(p, text, size=9)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        i += 1
    return rows, i


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Prospektimallin taulukuvaus")
    set_font(run, size=9, color="666666")


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_footer(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    title_done = False

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("!["):
            match = re.search(r"\(([^)]+)\)", stripped)
            if match:
                image_path = ROOT / match.group(1)
                if image_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(image_path), width=Inches(6.3))
            i += 1
            continue

        if stripped.startswith("|"):
            table_rows, i = parse_table(lines, i)
            add_table(doc, table_rows)
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not title_done:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(text)
                set_font(run, size=24, bold=True, color="0B2545")
                title_done = True
            else:
                p = doc.add_heading(level=1)
                add_markdown_text(p, text)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_markdown_text(p, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_markdown_text(p, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                add_markdown_text(p, lines[i].strip()[2:].strip())
                i += 1
            continue

        if re.match(r"\d+\.\s", stripped):
            while i < len(lines) and re.match(r"\d+\.\s", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                add_markdown_text(p, re.sub(r"^\d+\.\s*", "", lines[i].strip()))
                i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("![") or nxt.startswith("- ") or re.match(r"\d+\.\s", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_markdown_text(p, " ".join(paragraph_lines))

    doc.core_properties.title = "Prospektimallin taulukuvaus"
    doc.core_properties.subject = "Tietokantataulut prospektimallin automaattiajoon"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
